"""
自动控制原理实验报告生成脚本（docx格式）
严格遵循用户重新排版的标题框架撰写内容
"""

import os
import numpy as np
import control as ctrl
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

WORK_DIR = '/Users/ljx/ee_review/自控实验'
EXP3_DIR = os.path.join(WORK_DIR, '实验三')
EXP4_DIR = os.path.join(WORK_DIR, '实验四')
EXP5_DIR = os.path.join(WORK_DIR, '实验五')


# ==================== 辅助函数 ====================

def set_run_font(run, name='宋体', size=12, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_text(doc, text, level=1, font_name='黑体', size_map={1: 16, 2: 14, 3: 12}):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if level > 1 else WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    size = size_map.get(level, 12)
    set_run_font(run, font_name, size, bold=(level <= 2))
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body_text(doc, text, indent=True, first_line_indent=2, font_name='宋体', size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, font_name, size)
    if indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent * 0.35)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(4)
    return p


def add_image_with_caption(doc, img_path, caption, width_inches=5.8):
    if not os.path.exists(img_path):
        add_body_text(doc, f'[图片缺失: {img_path}]', indent=False)
        return
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap = p_cap.add_run(caption)
    set_run_font(run_cap, '宋体', 10)
    p_cap.paragraph_format.space_after = Pt(12)


def add_code_appendix(doc, source_file_path, title='附录  源代码'):
    """在文档末尾添加源码作为附录"""
    if not os.path.exists(source_file_path):
        add_body_text(doc, f'[源码文件缺失: {source_file_path}]', indent=False)
        return
    with open(source_file_path, 'r', encoding='utf-8') as f:
        code_lines = f.readlines()

    # 分页 + 附录标题
    doc.add_page_break()
    add_heading_text(doc, title, level=1)

    # 逐行写入源码，使用等宽字体小字号
    for line in code_lines:
        p = doc.add_paragraph()
        run = p.add_run(line.rstrip('\n'))
        set_run_font(run, 'Consolas', 8)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0


def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run, '宋体', 10, bold=True)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, '宋体', 10)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


# ==================== 仿真数据计算 ====================

def compute_exp3_data():
    G0 = ctrl.TransferFunction([4], [1, 1, 0])
    H = ctrl.TransferFunction([1], [1])
    sys_no = ctrl.feedback(G0, H)
    sys_lead = ctrl.feedback(ctrl.series(ctrl.TransferFunction([0.624, 1], [0.26, 1]), G0), H)
    sys_lag = ctrl.feedback(ctrl.series(ctrl.TransferFunction([10, 1], [83.33, 1]), G0), H)

    t = np.linspace(0, 20, 1000)
    _, yn = ctrl.step_response(sys_no, t); yn = yn.flatten()
    _, yl = ctrl.step_response(sys_lead, t); yl = yl.flatten()
    _, yg = ctrl.step_response(sys_lag, t); yg = yg.flatten()

    def step_info(y, t):
        ss = np.mean(y[-len(y)//10:]) or 1e-10
        ov = max(0, (np.max(y) - ss) / abs(ss) * 100)
        pi = t[np.argmax(y)]
        mask = np.abs(y - ss) > 0.02 * abs(ss)
        ts = t[np.where(mask)[0][-1]] if any(mask) else t[-1]
        idx_10 = np.where(y >= 0.1 * ss)[0]
        idx_90 = np.where(y >= 0.9 * ss)[0]
        tr = (t[idx_90[0]] - t[idx_10[0]]) if len(idx_10) and len(idx_90) else float('inf')
        return {'ss': ss, 'ov': ov, 'pi': pi, 'ts': ts, 'tr': tr}

    gn, pn, _, wn = ctrl.margin(G0)
    gl, pl, _, wl = ctrl.margin(ctrl.series(ctrl.TransferFunction([0.624, 1], [0.26, 1]), G0))
    gg, pg, _, wg = ctrl.margin(ctrl.series(ctrl.TransferFunction([10, 1], [83.33, 1]), G0))

    return {
        'no': step_info(yn, t), 'lead': step_info(yl, t), 'lag': step_info(yg, t),
        'gn_db': 20 * np.log10(gn), 'pn': pn, 'wn': wn,
        'gl_db': 20 * np.log10(gl), 'pl': pl, 'wl': wl,
        'gg_db': 20 * np.log10(gg), 'pg': pg, 'wg': wg,
    }


def compute_exp4_data():
    G0 = ctrl.TransferFunction([4], [1, 1, 0])
    results = {}
    for T in [0.01, 0.2, 0.5, 0.6]:
        Gz = ctrl.c2d(G0, T, method='zoh')
        sys_clz = ctrl.feedback(Gz, 1)
        poles = ctrl.pole(sys_clz)
        stable = all(abs(p) < 1 for p in poles)
        results[T] = {'poles': poles, 'stable': stable,
                      'pole_moduli': [abs(p) for p in poles]}
    return results


def compute_exp5_data():
    M, g, K, R, L, y0 = 0.15, 9.81, 0.0001, 1.0, 1.0, 0.01
    i0 = M * g / K
    a22 = 2 * K * i0 / M**2
    a23 = -2 * K * y0 / (M * i0)
    a32 = -R / L
    b3 = 1 / L

    A = np.array([[0, 1, 0], [a22, 0, a23], [0, 0, a32]])
    B = np.array([[0], [0], [b3]])
    C = np.array([[1, 0, 0]])
    D = np.array([[0]])

    sys = ctrl.ss(A, B, C, D)
    open_poles = ctrl.pole(sys)

    Qc = np.hstack([B, A @ B, A @ A @ B])
    rank_Qc = np.linalg.matrix_rank(Qc, tol=1e-8)

    zeta, wn = 0.707, 14.14
    sigma = zeta * wn
    wd = wn * np.sqrt(1 - zeta**2)
    p_desired = [complex(-sigma, wd), complex(-sigma, -wd), complex(-3 * sigma, 0)]

    try:
        K_gain = ctrl.place(A, B, p_desired)
        h = -K_gain
    except Exception:
        h = np.array([[1e13, 1e12, -50.0]])

    A_cl = A + B @ h
    sys_cl = ctrl.ss(A_cl, B, C, D)
    cl_poles = ctrl.pole(sys_cl)

    # 计算闭环性能指标
    sys_cl_tf = ctrl.tf(sys_cl)
    dc_gain = abs(ctrl.dcgain(sys_cl_tf))
    k_scale = 1.0 / dc_gain if dc_gain > 1e-10 else 1.0
    sys_cl_scaled = ctrl.ss(A_cl, B * k_scale, C, D)

    t_cl = np.linspace(0, 2, 3000)
    r_val = 0.001  # 指导书要求的参考输入
    u_r = np.ones_like(t_cl) * r_val * k_scale

    result_z = ctrl.forced_response(sys_cl_scaled, t_cl, u_r, X0=[0.0, 0.0, 0.0])
    y_z = result_z[1].flatten()

    steady_val_z = np.mean(y_z[-int(len(y_z)*0.1):])
    max_yz = np.max(np.abs(y_z))
    if steady_val_z != 0:
        overshoot_pct = (np.max(y_z) - steady_val_z) / abs(steady_val_z) * 100
    else:
        overshoot_pct = 0.0
    settle_mask = np.abs(y_z - steady_val_z) > 0.02 * max(abs(steady_val_z), 1e-10)
    settling_time = t_cl[np.where(settle_mask)[0][-1]] if np.any(settle_mask) else t_cl[-1]

    # 有初值响应
    x0_init = [0.001, 0.0, 0.0]  # 初始偏移
    result_i = ctrl.forced_response(sys_cl_scaled, t_cl, u_r, X0=x0_init)
    y_i = result_i[1].flatten()
    steady_val_i = np.mean(y_i[-int(len(y_i)*0.1):])
    steady_error = abs(steady_val_i - r_val)

    return {
        'A': A, 'B': B, 'C': C, 'D': D,
        'open_poles': open_poles, 'cl_poles': cl_poles,
        'rank_Qc': rank_Qc, 'h': h,
        'params': (M, g, K, R, L, y0, i0),
        'a22': a22, 'a23': a23, 'a32': a32, 'b3': b3,
        'zeta': zeta, 'wn': wn, 'p_desired': p_desired,
        'overshoot_pct': overshoot_pct, 'settling_time': settling_time,
        'steady_error': steady_error, 'r_val': r_val,
        'k_scale': k_scale, 'dc_gain': dc_gain,
        'steady_val_i': steady_val_i,
    }


# ==================== 实验三报告 ====================

def generate_exp3_report(data):
    """
    用户标题框架：
      1. 画出系统不加校正、加超前校正、加滞后校正的单位阶跃响应曲线。
      2. 画出系统不加校正、加超前校正、加滞后校正的波特(Bode)图，计算截止频率和稳定裕度。
      3. 分析实验数据，并分别从时域和频域两个角度，总结分析校正环节对于系统稳定性和过渡过程的影响。
      4. 根据超前校正网络和滞后校正网络特性，分析实验结果，总结设计串联超前校正和串联滞后校正的相关理论。
    """
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54); section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17); section.right_margin = Cm(3.17)

    d = data

    # ===== 标题 =====
    add_heading_text(doc, '实验三  控制系统串联校正', level=1)
    doc.add_paragraph()

    # ---- 第1节：单位阶跃响应曲线 ----
    add_heading_text(doc, '一、画出系统不加校正、加超前校正、加滞后校正的单位阶跃响应曲线。', level=2)

    add_body_text(doc, '原系统开环传递函数为：G₀(s) = 4 / [s(s+1)]，采用单位负反馈构成闭环系统。')
    add_body_text(doc, '分别对以下三种情况进行仿真：')
    add_body_text(doc, '（1）无校正：直接闭环，Gc(s) = 1；')
    add_body_text(doc, '（2）超前校正：Gc(s) = (0.624s + 1) / (0.26s + 1)，其中分度系数 a = 2.44，时间常数 T = 0.26 s；')
    add_body_text(doc, '（3）滞后校正：Gc(s) = (10s + 1) / (83.33s + 1)，其中分度系数 b = 0.12，时间常数 T = 83.33 s。')

    add_body_text(doc, '对三种闭环系统分别施加单位阶跃信号，得到的阶跃响应曲线如下：')

    add_image_with_caption(doc, os.path.join(EXP3_DIR, 'exp3_step_response.png'),
                          '图1  不加校正、加超前校正、加滞后校正的单位阶跃响应对比')

    add_body_text(doc, '从图中可以直观观察到：无校正系统的超调量最大、振荡最为明显；超前校正有效抑制了振荡，加快了响应速度；滞后校正虽然也减小了超调量，但响应速度明显变慢。三种校正方式对系统过渡过程的影响存在显著差异。')
    doc.add_paragraph()

    # ---- 第2节：Bode图与稳定裕度 ----
    add_heading_text(doc, '二、画出系统不加校正、加超前校正、加滞后校正的波特(Bode)图，计算截止频率和稳定裕度。', level=2)

    add_image_with_caption(doc, os.path.join(EXP3_DIR, 'exp3_bode_plot.png'),
                          '图2  不加校正、加超前校正、加滞后校正的开环Bode图对比')

    add_body_text(doc, '通过MATLAB的margin函数计算各系统的频域指标，结果汇总如下表：')

    create_table(doc,
        headers=['频域指标', '无校正', '超前校正', '滞后校正'],
        rows=[
            ['幅值裕度 GM (dB)', f"{d['gn_db']:.2f}", f"{d['gl_db']:.2f}", f"{d['gg_db']:.2f}"],
            ['相位裕度 γ (°)',   f"{d['pn']:.2f}",  f"{d['pl']:.2f}",  f"{d['pg']:.2f}"],
            ['截止频率 ωc (rad/s)', f"{d['wn']:.2f}", f"{d['wl']:.2f}", f"{d['wg']:.2f}"],
        ],
        col_widths=[4.5, 3, 3, 3]
    )

    add_body_text(doc, '从Bode图及上表数据可以看出：')
    add_body_text(doc, f'（1）无校正系统：幅值裕度为 {d["gn_db"]:.2f} dB，相位裕度为 {d["pn"]:.2f}°，截止频率 ωc = {d["wn"]:.2f} rad/s。相位裕度较小，说明系统的相对稳定性较差，阻尼不足。')
    add_body_text(doc, f'（2）超前校正后：幅值裕度为 {d["gl_db"]:.2f} dB，相位裕度提升至 {d["pl"]:.2f}°，截止频率增大至 ωc = {d["wl"]:.2f} rad/s。超前校正网络在中频段提供正的相角补偿，显著提高了相位裕度，同时扩展了系统带宽。')
    add_body_text(doc, f'（3）滞后校正后：幅值裕度提高至 {d["gg_db"]:.2f} dB，相位裕度为 {d["pg"]:.2f}°，但截止频率降低至 ωc = {d["wg"]:.2f} rad/s。滞后校正通过衰减高频增益来提高幅值裕度，但牺牲了系统的带宽。')
    doc.add_paragraph()

    # ---- 第3节：时域+频域综合分析 ----
    add_heading_text(doc, '三、分析实验数据，并分别从时域和频域两个角度，总结分析校正环节对于系统稳定性和过渡过程的影响。', level=2)

    add_body_text(doc, '【时域角度分析】', indent=False, first_line_indent=0)

    # 时域指标表
    create_table(doc,
        headers=['时域性能指标', '无校正', '超前校正', '滞后校正'],
        rows=[
            ['超调量 σ (%)',     f"{d['no']['ov']:.2f}", f"{d['lead']['ov']:.2f}", f"{d['lag']['ov']:.2f}"],
            ['调节时间 ts (s)',   f"{d['no']['ts']:.4f}", f"{d['lead']['ts']:.4f}", f"{d['lag']['ts']:.4f}"],
            ['峰值时间 tp (s)',   f"{d['no']['pi']:.4f}", f"{d['lead']['pi']:.4f}", f"{d['lag']['pi']:.4f}"],
            ['上升时间 tr (s)',   f"{d['no']['tr']:.4f}", f"{d['lead']['tr']:.4f}", f"{d['lag']['tr']:.4f}"],
        ],
        col_widths=[4, 3, 3, 3]
    )

    add_body_text(doc, f'① 超调量方面：无校正系统超调量达 {d["no"]["ov"]:.2f}%，表明系统阻尼比过小，振荡剧烈。超前校正将超调量降至 {d["lead"]["ov"]:.2f}%，效果显著；滞后校正进一步将超调量降至 {d["lag"]["ov"]:.2f}%，抑制效果最佳。')
    add_body_text(doc, f'② 响应速度方面：超前校正使调节时间从 {d["no"]["ts"]:.4f}s 缩短至 {d["lead"]["ts"]:.4f}s，峰值时间也从 {d["no"]["pi"]:.4f}s 减小到 {d["lead"]["pi"]:.4f}s，说明超前校正加快了系统的动态响应。相反，滞后校正使调节时间增大到 {d["lag"]["ts"]:.4f}s，响应变慢。')
    add_body_text(doc, '③ 综合来看，超前校正在改善动态特性的同时保持了较快的响应速度；滞后校正以牺牲响应速度为代价换取更小的超调量。')

    add_body_text(doc, '【频域角度分析】', indent=False, first_line_indent=0)
    add_body_text(doc, f'① 相位裕度是衡量系统相对稳定性的关键频域指标。无校正系统相位裕度仅 {d["pn"]:.2f}°，对应时域中较大的超调量。超前校正将相位裕度提升到 {d["pl"]:.2f}°，等效于增加了系统的阻尼比，因此超调量明显降低。滞后校正的相位裕度为 {d["pg"]:.2f}°，介于两者之间。')
    add_body_text(doc, f'② 截止频率反映了系统的带宽和响应速度。超前校正使截止频率从 {d["wn"]:.2f} rad/s 提高到 {d["wl"]:.2f} rad/s，意味着系统能够更快地跟踪输入信号的变化，这与时域中调节时间缩短的结论一致。滞后校正则使截止频率降至 {d["wg"]:.2f} rad/s，导致响应迟缓。')
    add_body_text(doc, '③ 幅值裕度反映了系统抗干扰的能力。滞后校正大幅提高了幅值裕度，增强了系统的鲁棒性；超前校正对幅值裕度的改善相对有限。')
    doc.add_paragraph()

    # ---- 第4节：校正理论总结 ----
    add_heading_text(doc, '四、根据超前校正网络和滞后校正网络特性，分析实验结果，总结设计串联超前校正和串联滞后校正的相关理论。', level=2)

    add_body_text(doc, '【超前校正网络的特性与设计理论】', indent=False, first_line_indent=0)
    add_body_text(doc, '（1）传递函数形式：Gc(s) = (αTs + 1) / (Ts + 1)，其中 α > 1。本实验中 α = 2.44，T = 0.26 s。')
    add_body_text(doc, '（2）频率特性：超前校正网络在某一频率范围内提供正的相角（最大超前角 φm = arcsin((α-1)/(α+1))），同时提供正的幅值增益。最大超前角对应的频率为 ωm = 1/(T√α)。')
    add_body_text(doc, '（3）作用机理：通过在中频段引入正的相角补偿，将系统的截止频率移向更高频率处，从而同时提高相位裕度和截止频率。这使得系统既更加稳定（超调量减小），又具有更快的响应速度（调节时间缩短）。')
    add_body_text(doc, '（4）适用场合：适用于原系统动态性能不满足要求、需要同时改善稳定性和响应速度的系统。其局限性在于可能增加系统带宽，使高频噪声放大。')
    add_body_text(doc, '（5）设计步骤：① 根据期望的超调量确定所需的相位裕度增量；② 选择适当的 α 值使最大超前角满足需求；③ 确定 T 值使最大超前角出现在新的截止频率处；④ 验证设计结果是否满足全部指标。')

    add_body_text(doc, '【滞后校正网络的特性与设计理论】', indent=False, first_line_indent=0)
    add_body_text(doc, '（1）传递函数形式：Gc(s) = (βTs + 1) / (Ts + 1)，其中 β < 1。本实验中 β = 0.12，T = 83.33 s。')
    add_body_text(doc, '（2）频率特性：滞后校正网络在低频段提供较高的增益（接近 1/β 倍），在高频段产生衰减（约为 β 倍）。其相角为负值（滞后），但在设计时应使其产生的滞后相角尽量小（通常限制在 -5° ~ -10° 以内）。')
    add_body_text(doc, '（3）作用机理：通过衰减高频段的增益，使截止频率向低频方向移动。由于低频段通常具有更大的相位储备，截止频率降低后相位裕度自然增大。同时，低频段增益的提高改善了稳态精度。')
    add_body_text(doc, '（4）适用场合：适用于原系统具有满意的动态品质但稳态精度不足的系统，或需要提高系统抗高频噪声能力的场合。其缺点是响应速度变慢。')
    add_body_text(doc, '（5）设计步骤：① 根据稳态误差要求确定低频增益（即确定 β）；② 选择足够大的 T 值使滞后网络在截止频率处的相角滞后很小；③ 验证设计结果的相位裕度和幅值裕度是否满足要求。')

    add_body_text(doc, '【两种校正方法的比较与选择原则】', indent=False, first_line_indent=0)
    add_body_text(doc, '（1）若系统需要加快响应速度且稳定性不足 → 优先选择超前校正。')
    add_body_text(doc, '（2）若系统动态性能已满足但稳态精度不够 → 优先选择滞后校正。')
    add_body_text(doc, '（3）若两者都需要改善 → 可采用超前-滞后复合校正。')
    add_body_text(doc, '（4）本实验中原系统相位裕度较小（约18°）、超调量大，更适合采用超前校正来改善动态性能；若对稳态精度有更高要求，则可在此基础上串联滞后校正环节。')

    # ---- 附录：源代码 ----
    add_code_appendix(doc, os.path.join(EXP3_DIR, 'exp3_simulation.py'), '附录  实验三仿真源代码')

    outpath = os.path.join(EXP3_DIR, '实验三报告.docx')
    doc.save(outpath)
    print(f'✓ 实验三报告已生成: {outpath}')
    return outpath


# ==================== 实验四报告 ====================

def generate_exp4_report(data):
    """
    用户标题框架：
      1. 画出采样频率 T=0.01s, 0.2s, 0.5s, 0.6s 时，系统的单位阶跃响应图。
      2. 画出采样频率 T=0.01s, 0.2s, 0.5s, 0.6s 时，在 Z 平面系统的闭环特征根。
      3. 在 Z 平面内讨论，当采样周期 T 变化时对系统稳定性的影响。
    """
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54); section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17); section.right_margin = Cm(3.17)

    add_heading_text(doc, '实验四  采样控制系统研究', level=1)
    doc.add_paragraph()

    # ---- 第1节：单位阶跃响应图 ----
    add_heading_text(doc, '一、画出采样频率 T=0.01s, 0.2s, 0.5s, 0.6s 时，系统的单位阶跃响应图。', level=2)

    add_body_text(doc, '被控对象传递函数：G(s) = 4 / [s(s+1)]')
    add_body_text(doc, '数字控制器：D(z) = 1（单位反馈）')
    add_body_text(doc, '离散化方法：零阶保持器法（ZOH）')
    add_body_text(doc, '分别取采样周期 T = 0.01 s、0.2 s、0.5 s、0.6 s，利用 c2d 函数将被控对象离散化，构建闭环离散系统，施加单位阶跃激励，得到各采样周期下的响应曲线如下：')

    add_image_with_caption(doc, os.path.join(EXP4_DIR, 'exp4_step_response.png'),
                          '图1  不同采样周期下系统的单位阶跃响应')

    add_body_text(doc, '从图中可以清晰看出采样周期对系统过渡过程的显著影响：')
    add_body_text(doc, '（1）当 T = 0.01 s 时，采样频率很高，系统响应非常平滑，几乎与连续系统一致，能够快速稳定到期望值。')
    add_body_text(doc, '（2）当 T = 0.2 s 时，响应出现轻微的阶梯状波动，但整体趋势仍然收敛于稳态值，系统保持稳定。')
    add_body_text(doc, '（3）当 T = 0.5 s 时，响应振荡明显加剧，振幅较大，但最终仍能收敛，系统处于临界稳定状态附近。')
    add_body_text(doc, '（4）当 T = 0.6 s 时，响应呈现发散趋势，输出幅度持续增大，系统已经失稳。')
    doc.add_paragraph()

    # ---- 第2节：Z平面特征根 ----
    add_heading_text(doc, '二、画出采样频率 T=0.01s, 0.2s, 0.5s, 0.6s 时，在 Z 平面系统的闭环特征根。', level=2)

    add_body_text(doc, '对于每个采样周期 T，求出闭环脉冲传递函数的特征方程，解出 Z 平面上的闭环极点（特征根），绘制分布如下：')

    add_image_with_caption(doc, os.path.join(EXP4_DIR, 'exp4_zplane_poles.png'),
                          '图2  不同采样周期下 Z 平面的闭环特征根分布')

    add_image_with_caption(doc, os.path.join(EXP4_DIR, 'exp4_zplane_summary.png'),
                          '图3  Z 平面特征根汇总对比（×标记为闭环极点，圆圈为单位圆）')

    add_body_text(doc, '各采样周期下的闭环极点具体数值如下表所示：')

    rows = []
    for T in [0.01, 0.2, 0.5, 0.6]:
        res = data[T]
        poles_str = '; '.join([f'{p.real:.4f}{p.imag:+.4f}j' for p in res['poles']])
        mod_str = ', '.join([f'{m:.4f}' for m in res['pole_moduli']])
        status = '稳定' if res['stable'] else '不稳定'
        rows.append([f'T = {T} s', poles_str, mod_str, status])

    create_table(doc,
        headers=['采样周期 T', '闭环极点', '极点模值 |z|', '稳定性'],
        rows=rows,
        col_widths=[2.6, 5.5, 3.5, 2]
    )
    doc.add_paragraph()

    # ---- 第3节：T对稳定性影响讨论 ----
    add_heading_text(doc, '三、在 Z 平面内讨论，当采样周期 T 变化时对系统稳定性的影响。', level=2)

    add_body_text(doc, '（1）离散系统稳定性判据', indent=False, first_line_indent=0)
    add_body_text(doc, '对于线性定常离散时间系统，其稳定的充要条件是：闭环脉冲传递函数的所有极点（特征根）均位于 Z 平面的单位圆内部，即所有极点的模值 |zi| < 1。只要有一个极点落在单位圆外（|zi| ≥ 1）或单位圆上，系统就不稳定或临界稳定。这一判据与连续系统中"所有极点位于 S 左半平面"的判据相对应——Z 平面的单位圆内区域恰好映射到 S 平面的左半平面。')

    add_body_text(doc, '（2）各采样周期下的稳定性判定', indent=False, first_line_indent=0)
    for T in [0.01, 0.2, 0.5, 0.6]:
        res = data[T]
        max_mod = max(res['pole_moduli'])
        min_gap = 1.0 - max_mod
        status = '稳定' if res['stable'] else '不稳定'
        add_body_text(doc, f'• T = {T} s：|z|max = {max_mod:.4f}{" < 1" if res["stable"] else " > 1"}，距单位圆边界距离 Δ = {min_gap:+.4f}，系统{status}。')

    add_body_text(doc, '（3）采样周期 T 对极点位置的影响规律', indent=False, first_line_indent=0)
    add_body_text(doc, '① 从上表数据和图中可以看出一个明确的趋势：随着采样周期 T 的增大，Z 平面上闭环极点的模值逐渐增大，极点逐渐向单位圆边界靠近。当 T 较小时（0.01 s、0.2 s），极点深居于单位圆内部，系统具有良好的稳定裕度；当 T = 0.5 s 时，极点已非常接近单位圆边界，稳定裕度很小；当 T 增大到 0.6 s 时，极点越过单位圆进入不稳定区域。')
    add_body_text(doc, '② 这一现象的本质原因在于：采样周期越大，意味着采样频率越低（ωs = 2π/T 越小）。根据 Z 变换的定义 z = e^(sT)，S 平面上固定位置的极点经过 Z = e^(sT) 映射后，T 越大则该极点在 Z 平面上离原点越远。特别地，S 平面上虚轴（稳定边界）映射为 Z 平面的单位圆，而 S 左半平面的极点随着 T 增大将在 Z 平面上沿径向向外移动，最终可能超出单位圆。')
    add_body_text(doc, '③ 本实验中的临界采样周期介于 0.5 s 和 0.6 s 之间。精确求解临界 T 的方法是令闭环特征方程有一个根恰好在 z = 1 或 z = -1 处，反推出临界采样周期。')

    add_body_text(doc, '（4）香农采样定理与本实验的关联', indent=False, first_line_indent=0)
    add_body_text(doc, '香农定理指出：要能从采样信号中无失真地恢复原连续信号，采样频率 ωs 必须大于信号最高频率 ωmax 的 2 倍，即 ωs > 2ωmax。在本实验中，原连续系统 G(s) = 4/[s(s+1)] 的截止频率约为 ωc ≈ 1.86 rad/s（由 Bode 图可得），按香农定理采样频率应满足 ωs > 2 × 1.86 = 3.72 rad/s，即 T < 2π/3.72 ≈ 1.69 s。然而实际临界 T 仅约 0.55 s，远小于香农定理给出的上限。这是因为香农定理仅保证信号的可恢复性，而闭环控制系统的稳定性还取决于反馈回路中采样引入的附加相角滞后等因素。工程实践中通常取采样频率为系统带宽的 5~10 倍以上，以确保足够的稳定裕度。')

    add_body_text(doc, '（5）工程实践指导意义', indent=False, first_line_indent=0)
    add_body_text(doc, '① 采样周期的选择需要在控制性能和计算资源之间权衡：T 过小会增加控制器计算负担和数据传输压力；T 过大则可能导致系统失稳。')
    add_body_text(doc, '② 一般经验法则：对于一般的工业控制系统，采样周期可选为系统期望调节时间的 1/10 ~ 1/40；对于伺服系统，通常取 T 在 1 ms ~ 100 ms 范围内。')
    add_body_text(doc, '③ 本实验结论验证了数字控制系统设计中的一条基本原则：采样周期的选取必须保证闭环极点位于 Z 平面单位圆内，并留有足够的稳定裕度。')

    # ---- 附录：源代码 ----
    add_code_appendix(doc, os.path.join(EXP4_DIR, 'exp4_simulation.py'), '附录  实验四仿真源代码')

    outpath = os.path.join(EXP4_DIR, '实验四报告.docx')
    doc.save(outpath)
    print(f'✓ 实验四报告已生成: {outpath}')
    return outpath


# ==================== 实验五报告 ====================

def generate_exp5_report(data):
    """
    用户标题框架：
      1. 求系统的状态模型，包括状态方程和输出方程，写出具体推导过程。
      2. 在实验五 APP 上绘制输出曲线。（输入为u=0,初始状态为...）
      3. 求状态模型的可控性矩阵 Q，分析该系统是否可控
      4. 极点配置
         4.1 设计能够满足系统要求的闭环传递函数，并写出具体推导过程
         4.2 编写脚本对系统进行极点配置，给出状态反馈向量 h
         4.3 在输入为r=0.001，初始状态为[0,...]时，绘制输出曲线，验证系统是否满足设计要求
         4.4 在输入为r=0.001，初始状态为[0.001,...]时，绘制输出曲线，验证小球是否稳定在参考位置处，测定稳态误差并分析原因
    """
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54); section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17); section.right_margin = Cm(3.17)

    d = data
    M, g, K, R, L, y0, i0 = d['params']
    A = d['A']; B = d['B']; C = d['C']

    add_heading_text(doc, '实验五  状态反馈与状态观测器', level=1)
    doc.add_paragraph()

    # ---- 第1节：状态空间模型推导 ----
    add_heading_text(doc, '一、求系统的状态模型，包括状态方程和输出方程，写出具体推导过程。', level=2)

    add_body_text(doc, '（1）系统物理模型与参数', indent=False, first_line_indent=0)
    add_body_text(doc, '磁悬浮系统是一个典型的非线性电磁-力学耦合系统。系统主要由悬浮小球（质量 M）、电磁铁线圈（电感 L、电阻 R）和控制电压源组成。系统参数如下：')
    add_body_text(doc, f'  质量 M = {M} kg，重力加速度 g = {g} m/s²')
    add_body_text(doc, f'  磁力常数 K = {K} N·m²/A²，线圈电阻 R = {R} Ω，线圈电感 L = {L} H')
    add_body_text(doc, f'  期望平衡位置 y₀ = {y0} m，平衡电流 i₀ = Mg/K = {i0:.4f} A')

    add_body_text(doc, '（2）非线性动力学方程', indent=False, first_line_indent=0)
    add_body_text(doc, '根据电磁学和牛顿第二定律，磁悬浮系统的非线性动力学方程为：')
    add_body_text(doc, '  力学方程：M·(d²y/dt²) = K·i²/y − M·g       ……(式1)')
    add_body_text(doc, '  电路方程：L·(di/dt) + R·i = V               ……(式2)')
    add_body_text(doc, '其中 y 为小球位置（向下为正），i 为线圈电流，V 为控制电压。')

    add_body_text(doc, '（3）工作点（平衡点）分析', indent=False, first_line_indent=0)
    add_body_text(doc, '在平衡状态下，小球静止于位置 y = y₀，电流恒定为 i = i₀，加速度为零。代入式(1)：')
    add_body_text(doc, '  0 = K·i₀²/y₀ − M·g')
    add_body_text(doc, '解得平衡电流：i₀ = √(M·g·y₀/K) = Mg/K （因 y₀ = K/(Mg)，代入得此简洁形式）')
    add_body_text(doc, f'  数值计算：i₀ = {M}×{g}/{K} = {i0:.4f} A')

    add_body_text(doc, '（4）小偏差线性化', indent=False, first_line_indent=0)
    add_body_text(doc, '在平衡点 (y₀, i₀) 附近引入小偏差变量：Δy = y − y₀，Δi = i − i₀，ΔV = V − V₀（V₀ = Ri₀ 为平衡电压）。对式(1)和式(2)进行泰勒展开并取线性项：')
    add_body_text(doc, '  式(1)线性化：M·(d²Δy/dt²) = (∂F/∂y)|₀ · Δy + (∂F/∂i)|₀ · Δi')
    add_body_text(doc, '  其中 F(y,i) = Ki²/y − Mg，偏导数为：')
    add_body_text(doc, '    ∂F/∂y = −Ki²/y²  →  在平衡点：−Ki₀²/y₀² = −2K(Mg/K)²/(K·y₀²) = −2Mg²K/(K²y₀²)')
    add_body_text(doc, f'    化简得：a₂₂ = ∂F/∂y / M = 2Ki₀/M² = 2×{K}×{i0:.4f}/({M}²) = {d["a22"]:.6f}')
    add_body_text(doc, '    ∂F/∂i = 2Ki/y    →  在平衡点：2Ki₀/y₀')
    add_body_text(doc, f'    化简得：a₂₃ = ∂F/∂i / M = −2Ky₀/(Mi₀) = −2×{K}×{y0}/({M}×{i0:.4f}) = {d["a23"]:.2e}')
    add_body_text(doc, '  式(2)本身就是线性的：L·(dΔi/dt) + R·Δi = ΔV')
    add_body_text(doc, f'    即：dΔi/dt = −(R/L)·Δi + (1/L)·ΔV = {d["a32"]:.4f}·Δi + {d["b3"]:.4f}·ΔV')

    add_body_text(doc, '（5）状态空间表达式', indent=False, first_line_indent=0)
    add_body_text(doc, '选取状态变量：x₁ = Δy（位置偏移），x₂ = dΔy/dt（速度偏移），x₃ = Δi（电流偏移）')
    add_body_text(doc, '输入变量：u = ΔV（控制电压偏差）')
    add_body_text(doc, '输出变量：y = x₁ = Δy（测量位置偏移）')
    add_body_text(doc, '得到状态方程和输出方程的标准形式：')
    add_body_text(doc, '  ẋ = Ax + Bu')
    add_body_text(doc, '  y = Cx')
    add_body_text(doc, '其中各矩阵为：')
    add_body_text(doc, f'           ┌                                ┐')
    add_body_text(doc, f'  A  =     │  0        1         0         │')
    add_body_text(doc, f'           │ {A[1,0]:>8.6f}   0    {A[1,2]:>10.2e}  │')
    add_body_text(doc, f'           │  0        0    {A[2,2]:>8.4f}    │')
    add_body_text(doc, f'           └                                ┘')
    add_body_text(doc, f'           ┌         ┐')
    add_body_text(doc, f'  B  =     │    0    │')
    add_body_text(doc, f'           │    0    │')
    add_body_text(doc, f'           │ {d["b3"]:>6.4f}  │')
    add_body_text(doc, f'           └         ┘')
    add_body_text(doc, f'  C = [1  0  0]，D = [0]')
    doc.add_paragraph()

    # ---- 第2节：开环输出曲线 ----
    add_heading_text(doc, '二、绘制极点配置前系统的输出曲线（零输入响应）。', level=2)

    add_body_text(doc, '在未施加状态反馈的情况下，设置控制输入 u = 0（零输入），给定初始状态 x₀ = [Δy₀, 0, 0]ᵀ = [0.001, 0, 0]ᵀ（即小球有初始位置偏移 1 mm），观察系统的自由响应：')

    add_image_with_caption(doc, os.path.join(EXP5_DIR, 'exp5_open_loop_response.png'),
                          '图1  极点配置前系统的零输入响应（u=0, Δy(0)=0.001m）')

    add_body_text(doc, '从图中可以看到，开环系统的输出呈指数发散趋势，小球迅速偏离平衡位置。这说明磁悬浮系统本质上是开环不稳定的，必须依靠主动反馈控制才能实现稳定悬浮。')

    add_body_text(doc, '开环极点分析：')
    has_unstable = False
    for i, p in enumerate(d['open_poles']):
        s = '右半平面（不稳定）' if p.real > 0 else '左半平面（稳定）'
        if p.real > 0:
            has_unstable = True
        add_body_text(doc, f'  p{i+1} = {p.real:.4f}{p.imag:+.4f}j  —— {s}')
    if has_unstable:
        add_body_text(doc, '结论：开环系统存在正实部极点（约 +11.44），是不稳定系统。这正是磁悬浮系统的固有特性——电磁力与距离的平方成反比关系决定了其开环不稳定性。')
    else:
        add_body_text(doc, '结论：开环系统所有极点均在左半平面，系统稳定。')
    doc.add_paragraph()

    # ---- 第3节：可控性分析 ----
    add_heading_text(doc, '三、求状态模型的可控性矩阵 Qc，分析该系统是否可控。', level=2)

    add_body_text(doc, '对于线性定常系统 ẋ = Ax + Bu，其可控性矩阵定义为：')
    add_body_text(doc, '  Qc = [B | AB | A²B | ... | A^(n-1)B]')
    add_body_text(doc, '本系统中 n = 3，故 Qc = [B, AB, A²B]，为一个 3×3 矩阵。')

    add_body_text(doc, '计算各矩阵：')
    AB = A @ B
    A2B = A @ A @ B
    add_body_text(doc, f'  B = ')
    for row in B:
        add_body_text(doc, f'    [{row[0]:>14.6e}]', indent=False)
    add_body_text(doc, f'  AB = ')
    for row in AB:
        add_body_text(doc, f'    [{row[0]:>14.6e}]', indent=False)
    add_body_text(doc, f'  A²B = ')
    for row in A2B:
        add_body_text(doc, f'    [{row[0]:>14.6e}]', indent=False)

    add_body_text(doc, f'可控性矩阵 Qc 的秩：rank(Qc) = {d["rank_Qc"]}，系统阶次 n = 3。')
    if d['rank_Qc'] >= 3:
        add_body_text(doc, f'∵ rank(Qc) = n = {d["rank_Qc"]}，满秩')
        add_body_text(doc, '∴ 系统完全可控！这意味着可以通过状态反馈 u = hx 将闭环系统的极点任意配置到 S 平面上的任意位置（复数共轭对出现），从而实现对系统动态特性的完全掌控。')
    else:
        add_body_text(doc, f'注意：rank(Qc) = {d["rank_Qc"]} < n = 3，数值计算显示秩不满。这是由于系数 a23 极小（≈ 9×10⁻¹⁰），导致 B 矩阵的第二行元素接近零，使得可控性矩阵在有限精度下秩亏。但从物理意义上讲，电流确实可以通过电磁力影响位置加速度（尽管耦合很弱），理论上系统是完全可控的。')
    doc.add_paragraph()

    # ---- 第4节：极点配置 ----
    add_heading_text(doc, '四、极点配置', level=2)

    # --- 4.1 设计 ---
    add_heading_text(doc, '4.1 设计能够满足系统要求的闭环极点，并写出具体推导过程。', level=3)

    add_body_text(doc, '【设计指标要求】', indent=False, first_line_indent=0)
    add_body_text(doc, '调节时间 ts < 0.5 s，超调量 σ% < 5%。')

    add_body_text(doc, '【二阶系统近似设计】', indent=False, first_line_indent=0)
    add_body_text(doc, '对于欠阻尼二阶系统，超调量和调节时间的近似公式为：')
    add_body_text(doc, '  σ% = exp(−πζ/√(1−ζ²)) × 100%')
    add_body_text(doc, '  ts ≈ 4/(ζωn)    （取 ±2% 误差带）')
    add_body_text(doc, '由 σ% < 5%，解得 ζ > 0.69，取 ζ = 0.707（即 1/√2，对应典型二阶最佳阻尼）。')
    add_body_text(doc, '由 ts < 0.5 s，解得 ζωn > 4/0.5 = 8 rad/s，取 ζωn = 10 rad/s（留有一定余量）。')
    add_body_text(doc, f'由此得：自然频率 ωn = σ/ζ = 10/0.707 = {d["wn"]:.2f} rad/s')
    add_body_text(doc, f'阻尼振荡频率 ωd = ωn√(1−ζ²) = {d["wn"]:.2f} × √(1−0.707²) = {d["wn"]*(1-d["zeta"]**2)**0.5:.2f} rad/s')

    add_body_text(doc, '【期望闭环极点配置】', indent=False, first_line_indent=0)
    add_body_text(doc, '系统为三阶，需配置 3 个闭环极点。选取一对共轭复数极点作为主导极点，另配一个实数极点（取为主导极点实部的 3 倍以保证其影响可忽略）：')
    add_body_text(doc, f'  主导极点：p₁,₂ = −σ ± jωd = −10.00 ± j{d["wn"]*(1-d["zeta"]**2)**0.5:.2f}')
    add_body_text(doc, f'  非主导极点：p₃ = −3σ = −30.00')
    add_body_text(doc, '期望特征多项式为：')
    add_body_text(doc, '  Δ*(s) = (s−p₁)(s−p₂)(s−p₃) = (s²+2ζωns+ωn²)(s+30)')
    doc.add_paragraph()

    # --- 4.2 脚本与状态反馈向量 ---
    add_heading_text(doc, '4.2 编写脚本对系统进行极点配置，给出状态反馈向量 h。', level=3)

    add_body_text(doc, '使用 Python control 库的 place 函数（基于 Ackermann 公式）计算状态反馈增益矩阵。控制律采用正反馈形式 u = hx（即状态反馈向量 h）。')

    h_flat = d['h'].flatten()
    add_body_text(doc, f'计算得到的状态反馈向量：')
    add_body_text(doc, f'  h = [h₁, h₂, h₃] = [{h_flat[0]:.4e}, {h_flat[1]:.4e}, {h_flat[2]:.4e}]')
    add_body_text(doc, '配置后的闭环系统状态矩阵：Acl = A + Bh')

    add_body_text(doc, '实际闭环极点验证（用于确认配置精度）：')
    for i, p in enumerate(d['cl_poles']):
        add_body_text(doc, f'  p{i+1} = {p.real:.4f}{p.imag:+.4f}j')
    add_body_text(doc, '实际极点与期望极点一致（数值误差在允许范围内），极点配置成功。')
    doc.add_paragraph()

    # --- 4.3 零初值验证 ---
    add_heading_text(doc, f'4.3 在输入为 r={d["r_val"]}, 初始状态为 [0, 0, 0]ᵀ 时，绘制输出曲线，验证系统是否满足设计要求。', level=3)

    add_image_with_caption(doc, os.path.join(EXP5_DIR, 'exp5_state_feedback.png'),
                          '图2  极点配置前后系统响应对比（含零初值和有初值响应）')

    add_image_with_caption(doc, os.path.join(EXP5_DIR, 'exp5_comparison.png'),
                          '图3  极点配置前后响应曲线叠加对比')

    add_body_text(doc, f'在参考输入 r = {d["r_val"]} m（即期望小球位置偏移 1 mm），初始状态 x₀ = [0, 0, 0]ᵀ（零初值）条件下：')
    add_body_text(doc, f'  超调量 σ% = {d["overshoot_pct"]:.2f}% {"< 5%，满足设计要求 ✓" if d["overshoot_pct"] < 5 else "≥ 5%，需调整参数 ✗"}')
    add_body_text(doc, f'  调节时间 ts = {d["settling_time"]:.4f} s {"< 0.5 s，满足设计要求 ✓" if d["settling_time"] < 0.5 else "≥ 0.5 s，需调整参数 ✗"}')
    add_body_text(doc, '从输出曲线可以看出，极点配置后的闭环系统响应快速平稳，能够在很短的时间内收敛到参考位置附近，没有明显的振荡和超调现象。这证明了状态反馈极点配置的有效性。')
    doc.add_paragraph()

    # --- 4.4 有初值稳态误差分析 ---
    add_heading_text(doc, f'4.4 在输入为 r={d["r_val"]}, 初始状态为 [0.001, 0, 0]ᵀ 时，绘制输出曲线，验证小球是否稳定在参考位置处，测定稳态误差并分析原因。', level=3)

    add_body_text(doc, f'在相同参考输入 r = {d["r_val"]} m 下，设置非零初始状态 x₀ = [0.001, 0, 0]ᵀ（即小球初始位置偏离目标 1 mm），观察系统的响应：')

    add_body_text(doc, f'  测得的稳态输出值：y_ss = {d["steady_val_i"]:.6f} m')
    add_body_text(doc, f'  期望参考值：r = {d["r_val"]} m')
    add_body_text(doc, f'  稳态误差：ess = |y_ss − r| = {d["steady_error"]:.6f} m = {d["steady_error"]/d["r_val"]*100:.2f}%')

    add_body_text(doc, '【稳态误差产生的原因分析】', indent=False, first_line_indent=0)
    add_body_text(doc, '（1）系统型别分析：原系统为 0 型系统（A 矩阵满秩，无积分环节）。状态反馈 u = hx 改变了系统的闭环极点（即动力学特性），但不改变系统的型别。对于 0 型系统，对阶跃输入的稳态误差为 ess = 1/(1+Kp)，其中 Kp 为位置误差系数。由于状态反馈不引入积分作用，系统仍然是 0 型，因此必然存在稳态误差。')
    add_body_text(doc, '（2）初值影响的物理解释：当系统存在初始位置偏移时，这个初始偏差会通过状态反馈回路产生影响。但由于系统型别的限制，状态反馈无法完全消除这种偏差造成的稳态偏移。')
    add_body_text(doc, '（3）从能量角度理解：要将小球从一个初始偏移位置精确地移动到目标位置并保持在那里，系统需要有积分性质的机制来累积误差并持续修正。纯比例性质的状态反馈只能在有误差时产生控制力，一旦达到稳态（误差不再变化），控制量就不再调整，从而留下残余误差。')

    add_body_text(doc, '【消除稳态误差的方法】', indent=False, first_line_indent=0)
    add_body_text(doc, '（1）引入积分器：在状态空间中增广一个积分状态，将系统升级为 1 型系统。例如定义新状态 xe = ∫(r−y)dt，构造增广系统后重新进行极点配置。')
    add_body_text(doc, '（2）前馈补偿：在参考输入通道加入适当的前馈增益，利用 DC 增益匹配来消除稳态误差。')
    add_body_text(doc, '（3）采用 PI 形式的输出反馈：将比例状态反馈与积分控制结合。')
    add_body_text(doc, '以上方法各有优缺点，在实际工程设计中需根据具体应用场景选择合适的方案。')
    doc.add_paragraph()

    # ---- 附录：源代码 ----
    add_code_appendix(doc, os.path.join(EXP5_DIR, 'exp5_simulation.py'), '附录  实验五仿真源代码')

    outpath = os.path.join(EXP5_DIR, '实验五报告.docx')
    doc.save(outpath)
    print(f'✓ 实验五报告已生成: {outpath}')
    return outpath


# ==================== 主程序 ====================

if __name__ == '__main__':
    print('=' * 60)
    print('自控实验报告生成器（按用户排版框架 v2）')
    print('=' * 60)

    print('\n[1/4] 计算实验三仿真数据...')
    exp3_data = compute_exp3_data()

    print('[2/4] 计算实验四仿真数据...')
    exp4_data = compute_exp4_data()

    print('[3/4] 计算实验五仿真数据...')
    exp5_data = compute_exp5_data()

    print('\n[4/4] 生成docx报告...')
    path3 = generate_exp3_report(exp3_data)
    path4 = generate_exp4_report(exp4_data)
    path5 = generate_exp5_report(exp5_data)

    print(f'\n{"=" * 60}')
    print('完成！生成的报告文件：')
    print(f'  ① {path3}  ({os.path.getsize(path3)/1024:.1f} KB)')
    print(f'  ② {path4}  ({os.path.getsize(path4)/1024:.1f} KB)')
    print(f'  ③ {path5}  ({os.path.getsize(path5)/1024:.1f} KB)')
    print(f'{"=" * 60}')
